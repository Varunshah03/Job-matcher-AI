import io
import logging
from typing import Optional
import aiofiles
import tempfile
import os
import re

# PDF processing
import pdfplumber

# DOCX processing
from docx import Document

# OCR for scanned documents
try:
    from paddleocr import PaddleOCR
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("PaddleOCR not available. OCR functionality will be disabled.")

from PIL import Image

# Clear existing handlers to prevent conflicts
for handler in logging.root.handlers[:]:
    logging.root.removeHandler(handler)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

class ResumeProcessor:
    """Handles resume text extraction from various file formats"""
    
    def __init__(self):
        logger.info(f"Initializing ResumeProcessor, OCR available: {OCR_AVAILABLE}")
        self.ocr = None
        if OCR_AVAILABLE:
            try:
                self.ocr = PaddleOCR(use_angle_cls=True, lang='en')
                logger.info("PaddleOCR initialized successfully")
            except Exception as e:
                logger.warning(f"Failed to initialize OCR: {str(e)}")
                self.ocr = None
    
    async def extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from uploaded file based on its type"""
        logger.info(f"Extracting text from {filename}")
        try:
            file_extension = filename.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                return await self._extract_from_pdf(file_content, filename)
            elif file_extension == 'docx':
                return await self._extract_from_docx(file_content, filename)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise
    
    async def _extract_from_pdf(self, pdf_content: bytes, filename: str) -> str:
        """Extract text from PDF file"""
        logger.info(f"Processing PDF: {filename}")
        try:
            text_content = []
            
            with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    logger.info(f"Extracting text from page {page_num + 1}")
                    # Try to extract text normally
                    text = page.extract_text()
                    
                    if text and text.strip():
                        logger.info(f"Extracted {len(text)} characters from page {page_num + 1}")
                        text_content.append(text)
                    elif self.ocr:
                        # Try OCR on the page image
                        logger.info(f"Attempting OCR on page {page_num + 1}")
                        try:
                            # Convert page to image
                            page_image = page.to_image(resolution=150)
                            
                            # Use OCR to extract text
                            ocr_result = self.ocr.ocr(page_image.original, cls=True)
                            
                            page_text = []
                            for line in ocr_result:
                                if line:
                                    for word_info in line:
                                        if word_info[1][0]:  # Text content
                                            page_text.append(word_info[1][0])
                            
                            if page_text:
                                ocr_text = ' '.join(page_text)
                                logger.info(f"OCR extracted {len(ocr_text)} characters from page {page_num + 1}")
                                text_content.append(ocr_text)
                            else:
                                logger.warning(f"No text extracted via OCR from page {page_num + 1}")
                                
                        except Exception as ocr_error:
                            logger.warning(f"OCR failed for page {page_num + 1}: {ocr_error}")
                    else:
                        logger.warning(f"No text extracted from page {page_num + 1} and OCR not available")
            
            extracted_text = '\n'.join(text_content)
            
            if not extracted_text.strip():
                logger.error(f"No readable text found in PDF: {filename}")
                raise ValueError("No readable text found in PDF")
            
            cleaned_text = self._clean_text(extracted_text)
            logger.info(f"Cleaned text for {filename}: {len(cleaned_text)} characters")
            return cleaned_text
            
        except Exception as e:
            logger.error(f"PDF extraction failed for {filename}: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_from_docx(self, docx_content: bytes, filename: str) -> str:
        """Extract text from DOCX file"""
        logger.info(f"Processing DOCX: {filename}")
        try:
            # Create a temporary file to work with python-docx
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(docx_content)
                temp_file_path = temp_file.name
            
            try:
                doc = Document(temp_file_path)
                text_content = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_content.append(cell.text)
                
                extracted_text = '\n'.join(text_content)
                
                if not extracted_text.strip():
                    logger.error(f"No readable text found in DOCX: {filename}")
                    raise ValueError("No readable text found in DOCX")
                
                cleaned_text = self._clean_text(extracted_text)
                logger.info(f"Cleaned text for {filename}: {len(cleaned_text)} characters")
                return cleaned_text
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except Exception:
                    pass
                    
        except Exception as e:
            logger.error(f"DOCX extraction failed for {filename}: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Join lines and normalize spaces
        cleaned_text = ' '.join(lines)
        
        # Remove multiple spaces
        import re
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        
        return cleaned_text.strip()
    
    def validate_extracted_text(self, text: str) -> bool:
        """Validate if extracted text is meaningful"""
        if not text or len(text.strip()) < 50:
            return False
        
        # Check if text contains some common resume keywords
        resume_keywords = [
            'experience', 'education', 'skill', 'work', 'project',
            'university', 'college', 'developer', 'engineer', 'manager',
            'bachelor', 'master', 'degree', 'certificate', 'programming',
            'marketing', 'digital', 'analytics', 'seo', 'content'
        ]
        
        text_lower = text.lower()
        keyword_count = sum(1 for keyword in resume_keywords if keyword in text_lower)
        
        # Should contain at least 2 common resume keywords
        return keyword_count >= 2